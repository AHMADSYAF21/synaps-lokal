"""
RAG Engine — Retrieval-Augmented Generation
Upload dokumen (PDF, TXT, MD, DOCX, CSV) → chunk → embed → simpan di ChromaDB
Query → semantic search → inject konteks → LLM menjawab berdasarkan dokumen
"""

import asyncio
import hashlib
import json
import logging
import re
import time
import uuid
from pathlib import Path
from typing import AsyncGenerator, Dict, List, Optional, Tuple

log = logging.getLogger("synapse.rag")

DOCS_DIR = Path("./data/documents")
DOCS_DIR.mkdir(parents=True, exist_ok=True)

CHUNK_SIZE   = 800   # characters per chunk
CHUNK_OVERLAP = 100  # overlap between chunks

RAG_SYSTEM = """You are a document analyst. Answer questions ONLY using the provided document excerpts.
If the answer is not in the documents, say so clearly.
Always cite which document/section you're referencing.
Be precise and factual."""


class Document:
    def __init__(self, doc_id: str, name: str, file_type: str,
                 chunk_count: int, size: int, created_at: float):
        self.doc_id      = doc_id
        self.name        = name
        self.file_type   = file_type
        self.chunk_count = chunk_count
        self.size        = size
        self.created_at  = created_at


class RAGEngine:
    def __init__(self, llm, memory):
        self.llm    = llm
        self.memory = memory
        self._docs: Dict[str, Document] = {}
        self._index_file = DOCS_DIR / "index.json"
        self._load_index()

    # ── Index Persistence ─────────────────────────────────────────────────────
    def _load_index(self):
        if self._index_file.exists():
            try:
                data = json.loads(self._index_file.read_text())
                for d in data:
                    self._docs[d["doc_id"]] = Document(**d)
                log.info(f"RAG: loaded {len(self._docs)} document(s)")
            except Exception as e:
                log.warning(f"RAG index load error: {e}")

    def _save_index(self):
        data = [vars(d) for d in self._docs.values()]
        self._index_file.write_text(json.dumps(data, indent=2))

    # ── Document Ingestion ────────────────────────────────────────────────────
    async def ingest(self, file_bytes: bytes, filename: str) -> Dict:
        """Parse, chunk, embed, and index a document."""
        ext = Path(filename).suffix.lower()
        doc_id = hashlib.md5(file_bytes).hexdigest()[:12]

        # Check if already indexed
        if doc_id in self._docs:
            return {"success": True, "doc_id": doc_id,
                    "message": "Document already indexed",
                    "chunks": self._docs[doc_id].chunk_count}

        # Save raw file
        raw_path = DOCS_DIR / f"{doc_id}{ext}"
        raw_path.write_bytes(file_bytes)

        # Extract text
        text, file_type = await self._extract_text(file_bytes, filename, ext)
        if not text:
            raw_path.unlink(missing_ok=True)
            return {"success": False, "error": "Could not extract text from document"}

        # Chunk text
        chunks = self._chunk_text(text, filename)
        if not chunks:
            return {"success": False, "error": "Document has no readable content"}

        # Save chunks to vector memory
        log.info(f"Indexing {len(chunks)} chunks from '{filename}'…")
        for i, chunk in enumerate(chunks):
            await self.memory.save(
                text=chunk,
                metadata={
                    "type":     "rag_chunk",
                    "doc_id":   doc_id,
                    "filename": filename,
                    "chunk_i":  i,
                    "total":    len(chunks),
                },
                collection="knowledge",
            )

        # Register document
        doc = Document(
            doc_id=doc_id, name=filename, file_type=file_type,
            chunk_count=len(chunks), size=len(file_bytes),
            created_at=time.time(),
        )
        self._docs[doc_id] = doc
        self._save_index()

        log.info(f"✅ Indexed '{filename}': {len(chunks)} chunks")
        return {"success": True, "doc_id": doc_id,
                "filename": filename, "chunks": len(chunks),
                "size_kb": round(len(file_bytes) / 1024, 1)}

    # ── Query ─────────────────────────────────────────────────────────────────
    async def query(self, question: str, doc_id: Optional[str] = None,
                    n_chunks: int = 6) -> Dict:
        """Search relevant chunks and answer using LLM."""
        chunks = await self._retrieve(question, doc_id, n_chunks)
        if not chunks:
            return {"answer": "No relevant content found in indexed documents.",
                    "sources": [], "context_used": 0}

        context = self._format_context(chunks)
        prompt  = f"{context}\n\nQuestion: {question}\n\nAnswer:"
        answer  = await self.llm.complete(prompt, role="general", system=RAG_SYSTEM)

        sources = list({c["metadata"]["filename"] for c in chunks})
        return {"answer": answer, "sources": sources,
                "context_used": len(chunks),
                "chunks": [{"text": c["text"][:200], "file": c["metadata"]["filename"],
                             "relevance": c["relevance"]} for c in chunks]}

    async def query_stream(self, question: str, doc_id: Optional[str] = None,
                           n_chunks: int = 6) -> AsyncGenerator[str, None]:
        chunks = await self._retrieve(question, doc_id, n_chunks)
        if not chunks:
            yield "No relevant content found in indexed documents."
            return

        context = self._format_context(chunks)
        sources = list({c["metadata"]["filename"] for c in chunks})

        yield f"[SOURCES: {', '.join(sources)}]\n\n"
        async for token in self.llm.stream(
            f"{context}\n\nQuestion: {question}\n\nAnswer:",
            role="general", system=RAG_SYSTEM
        ):
            yield token

    # ── Summarise entire document ─────────────────────────────────────────────
    async def summarise(self, doc_id: str) -> Dict:
        if doc_id not in self._docs:
            return {"success": False, "error": "Document not found"}

        chunks = await self._retrieve("", doc_id=doc_id, n=30)
        if not chunks:
            return {"success": False, "error": "No chunks found"}

        full_text = " ".join(c["text"] for c in chunks[:20])
        prompt    = f"Summarise this document concisely:\n\n{full_text[:6000]}"
        summary   = await self.llm.complete(prompt, role="general")
        return {"success": True, "doc_id": doc_id,
                "filename": self._docs[doc_id].name, "summary": summary}

    # ── List / Delete ─────────────────────────────────────────────────────────
    def list_docs(self) -> List[Dict]:
        return [{"doc_id": d.doc_id, "name": d.name, "file_type": d.file_type,
                 "chunks": d.chunk_count,
                 "size_kb": round(d.size / 1024, 1),
                 "created_at": d.created_at}
                for d in sorted(self._docs.values(), key=lambda x: -x.created_at)]

    async def delete_doc(self, doc_id: str) -> bool:
        if doc_id not in self._docs:
            return False
        # Remove raw file
        for f in DOCS_DIR.glob(f"{doc_id}.*"):
            f.unlink(missing_ok=True)
        # TODO: remove chunks from ChromaDB (would need direct collection access)
        del self._docs[doc_id]
        self._save_index()
        return True

    def stats(self) -> Dict:
        return {"total_documents": len(self._docs),
                "total_chunks": sum(d.chunk_count for d in self._docs.values()),
                "total_size_kb": round(sum(d.size for d in self._docs.values()) / 1024, 1)}

    # ── Private helpers ───────────────────────────────────────────────────────
    async def _extract_text(self, data: bytes, name: str, ext: str) -> Tuple[str, str]:
        loop = asyncio.get_running_loop()

        if ext in (".txt", ".md", ".rst", ".log"):
            for enc in ("utf-8", "latin-1", "cp1252"):
                try:
                    return data.decode(enc), "text"
                except UnicodeDecodeError:
                    continue
            return "", "text"

        elif ext == ".pdf":
            return await loop.run_in_executor(None, lambda: self._extract_pdf(data))

        elif ext in (".docx", ".doc"):
            return await loop.run_in_executor(None, lambda: self._extract_docx(data))

        elif ext == ".csv":
            try:
                import csv, io
                reader = csv.reader(io.StringIO(data.decode("utf-8", errors="replace")))
                rows   = [", ".join(row) for row in reader]
                return "\n".join(rows), "csv"
            except Exception as e:
                return "", "csv"

        elif ext == ".json":
            try:
                obj = json.loads(data)
                return json.dumps(obj, indent=2, ensure_ascii=False), "json"
            except Exception:
                return data.decode("utf-8", errors="replace"), "json"

        else:
            # Attempt raw text
            return data.decode("utf-8", errors="replace"), "unknown"

    def _extract_pdf(self, data: bytes) -> Tuple[str, str]:
        try:
            import pdfplumber, io
            text_parts = []
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
            return "\n\n".join(text_parts), "pdf"
        except ImportError:
            pass
        try:
            import PyPDF2, io
            reader = PyPDF2.PdfReader(io.BytesIO(data))
            pages  = [p.extract_text() or "" for p in reader.pages]
            return "\n\n".join(pages), "pdf"
        except Exception as e:
            log.warning(f"PDF extract error: {e}")
            return "", "pdf"

    def _extract_docx(self, data: bytes) -> Tuple[str, str]:
        try:
            import docx, io
            doc   = docx.Document(io.BytesIO(data))
            lines = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(lines), "docx"
        except Exception as e:
            log.warning(f"DOCX extract error: {e}")
            return "", "docx"

    def _chunk_text(self, text: str, source: str) -> List[str]:
        # Clean
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        if not text:
            return []

        chunks, start = [], 0
        while start < len(text):
            end   = min(start + CHUNK_SIZE, len(text))
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(f"[{source}]\n{chunk}")
            start = end - CHUNK_OVERLAP if end < len(text) else len(text)
        return chunks

    async def _retrieve(self, query: str, doc_id: Optional[str] = None,
                        n: int = 6) -> List[Dict]:
        try:
            results = await self.memory.search(
                query if query else "summary overview",
                n=n * 2,
                collection="knowledge",
            )
        except Exception:
            return []

        chunks = [r for r in results
                  if r.get("metadata", {}).get("type") == "rag_chunk"]
        if doc_id:
            chunks = [c for c in chunks
                      if c.get("metadata", {}).get("doc_id") == doc_id]
        return chunks[:n]

    def _format_context(self, chunks: List[Dict]) -> str:
        parts = ["[DOCUMENT EXCERPTS]"]
        for i, c in enumerate(chunks, 1):
            fname = c.get("metadata", {}).get("filename", "?")
            parts.append(f"\n--- Excerpt {i} (from: {fname}) ---\n{c['text']}")
        return "\n".join(parts)
